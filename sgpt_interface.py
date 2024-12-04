import streamlit as st
import logging
import requests
import json
import subprocess
import os
import sys
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from db_manager import DatabaseManager
import base64
from io import BytesIO
import pyperclip
from openai import OpenAI
import magic
import hashlib
import time
import pathlib
from pathlib import Path
import docx
from pypdf import PdfReader
from bs4 import BeautifulSoup
from workspace_manager import WorkspaceManager
from functools import lru_cache
from email_manager import EmailManager

# Initialize database manager
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "conversations.db")
db = DatabaseManager(DB_PATH)

# Print the current directory and file location for debugging
print(f"Current working directory: {os.getcwd()}")
print(f"Script location: {__file__}")

# Configure logging with absolute paths
APP_DIR = "/Users/christopher.bradford/mac_computer_use/local_shellgpt_llm"
log_dir = pathlib.Path(APP_DIR) / "logs"
log_dir.mkdir(exist_ok=True)
log_file = log_dir / f"sgpt_{datetime.now().strftime('%Y%m%d')}.log"

# Configure logging to both file and console
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(str(log_file)),
        logging.StreamHandler()
    ]
)

# Log initial startup information
logging.info("=== Application Starting ===")
logging.info(f"Working Directory: {os.getcwd()}")
logging.info(f"Script Location: {__file__}")
logging.info(f"Log File: {log_file}")

def load_api_key():
    """Load OpenAI API key from config file"""
    # First check environment variable
    api_key = os.environ.get('OPENAI_API_KEY')
    if api_key:
        return api_key

    # Then check config file
    config_file = pathlib.Path.home() / '.config' / 'shell-gpt' / '.sgptrc'
    if config_file.exists():
        try:
            content = config_file.read_text()
            if 'OPENAI_API_KEY=' in content:
                api_key = content.split('OPENAI_API_KEY=')[1].strip()
                os.environ['OPENAI_API_KEY'] = api_key
                return api_key
        except Exception as e:
            logging.error(f"Error loading API key: {str(e)}")
    return ''

def save_default_model(model_name):
    """Save the default model preference"""
    try:
        config_dir = os.path.expanduser("~/.config/shell-gpt")
        os.makedirs(config_dir, exist_ok=True)
        config_file = os.path.join(config_dir, "config.json")
        
        # Load existing config or create new
        config = {}
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                config = json.load(f)
        
        config['default_model'] = model_name
        
        with open(config_file, 'w') as f:
            json.dump(config, f)
            
        logging.info(f"Saved {model_name} as default model")
        return True
    except Exception as e:
        logging.error(f"Error saving default model: {str(e)}")
        return False

def load_default_model():
    """Load the default model preference"""
    try:
        config_file = os.path.expanduser("~/.config/shell-gpt/config.json")
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                config = json.load(f)
                return config.get('default_model')
    except Exception as e:
        logging.error(f"Error loading default model: {str(e)}")
    return None

def save_api_key(api_key):
    config_dir = pathlib.Path.home() / '.config' / 'shell-gpt'
    config_dir.mkdir(parents=True, exist_ok=True)
    config_file = config_dir / '.sgptrc'
    
    try:
        config_file.write_text(f'OPENAI_API_KEY={api_key}\n')
        os.environ['OPENAI_API_KEY'] = api_key
        return True
    except Exception as e:
        logging.error(f"Error saving API key: {str(e)}")
        return False

def execute_python_script(script):
    """Execute a Python script and return its output"""
    try:
        # Create a temporary directory for script execution
        tmp_dir = pathlib.Path.home() / '.sgpt_tmp'
        tmp_dir.mkdir(exist_ok=True)
        
        # Save the script to a temporary file
        script_file = tmp_dir / f"script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
        script_file.write_text(script)
        
        # Install required packages
        if 'pptx' in script:
            subprocess.run([sys.executable, "-m", "pip", "install", "python-pptx"], check=True)
        
        # Execute the script
        result = subprocess.run(
            [sys.executable, str(script_file)],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        # Clean up
        script_file.unlink()
        
        output = result.stdout.strip()
        if result.returncode != 0:
            output += f"\nError: {result.stderr.strip()}"
        return output
    except Exception as e:
        return f"Error executing Python script: {str(e)}"

def execute_shell_command(command):
    """Execute a shell command and return its output"""
    try:
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=30
        )
        return result.stdout.strip() if result.stdout else result.stderr.strip()
    except Exception as e:
        return f"Error executing command: {str(e)}"

def execute_web_search(query):
    """Execute a web search using DuckDuckGo and return the results"""
    try:
        logging.info(f"Performing web search for: {query}")
        url = f"https://duckduckgo.com/html/?q={query}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        results = []
        
        # Find all search result elements
        search_results = soup.find_all('div', class_='result')[:10]
        
        for result in search_results:
            try:
                title_elem = result.find('a', class_='result__a')
                snippet_elem = result.find('a', class_='result__snippet')
                
                if title_elem and snippet_elem:
                    title = title_elem.get_text(strip=True)
                    snippet = snippet_elem.get_text(strip=True)
                    link = title_elem.get('href')
                    
                    results.append({
                        "title": title,
                        "snippet": snippet,
                        "link": link
                    })
                    logging.info(f"Found result: {title}")
            except Exception as e:
                logging.warning(f"Error processing search result: {str(e)}")
                continue
        
        if not results:
            # Fallback in case no results were found
            results.append({
                "title": "No detailed results found",
                "snippet": f"Please try rephrasing your search query: {query}",
                "link": url
            })
        
        logging.info(f"Found {len(results)} search results")
        return results
    except Exception as e:
        logging.error(f"Error in web search: {str(e)}")
        raise Exception(f"Error executing web search: {str(e)}")

def call_ollama_api(model, prompt, stream=True):
    """Call Ollama API directly instead of using command line"""
    try:
        url = "http://localhost:11434/api/generate"
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": stream
        }
        
        if stream:
            response = requests.post(url, json=payload, stream=True)
            response.raise_for_status()
            
            full_response = []
            placeholder = st.empty()
            
            for line in response.iter_lines():
                if line:
                    json_response = json.loads(line)
                    if 'response' in json_response:
                        full_response.append(json_response['response'])
                        placeholder.markdown(''.join(full_response))
            
            final_response = ''.join(full_response)
            placeholder.empty()
            return final_response
        else:
            response = requests.post(url, json=payload)
            response.raise_for_status()
            result = response.json()
            return result.get('response', 'No response from model')
        
    except requests.exceptions.RequestException as e:
        logging.error(f"Ollama API error: {str(e)}")
        return f"Error calling Ollama API: {str(e)}"

def run_sgpt_command(prompt, model="gpt-3.5-turbo", use_ollama=False):
    # Remove 'sgpt' from the beginning of the prompt if user included it
    prompt = prompt.strip()
    if prompt.lower().startswith('sgpt'):
        prompt = prompt[4:].strip()
    
    logging.info(f"Processing prompt: {prompt}")
    logging.info(f"Model: {model}, Using Ollama: {use_ollama}")
    
    try:
        if use_ollama:
            import requests
            import json
            
            # Check if Ollama is available by testing the connection
            try:
                response = requests.get("http://localhost:11434/api/tags")
                if response.status_code != 200:
                    return "Error: Ollama server is not responding correctly. Please check if Ollama is running."
            except requests.exceptions.ConnectionError:
                logging.error("Ollama server not running")
                return "Error: Ollama server is not running. Please start Ollama first."
            
            # First, ask Ollama if this is a command to execute
            system_prompt = """You are a helpful AI assistant with the ability to execute commands and create files. When the user requests file operations or system commands, provide the appropriate command or Python script to execute.

For simple commands, respond with:
EXECUTE: <command>

For complex operations requiring Python, respond with:
PYTHON:
<complete python script>
END_PYTHON

Examples:
User: "list files in current directory"
Response: EXECUTE: ls -la

User: "Who is the president of France?"
Response: Emmanuel Macron is the current President of France...

User: "Create a text file with a poem about cats"
Response: PYTHON:
with open('cat_poem.txt', 'w') as f:
    f.write('''Whiskers twitching in the night,
Paws so soft and eyes so bright.
Purring softly on my lap,
Taking yet another nap.''')
END_PYTHON

Start your response with either 'EXECUTE: ' for simple commands or 'PYTHON:' for complex operations."""

            try:
                response = requests.post(
                    "http://localhost:11434/api/chat",
                    json={
                        "model": model,
                        "messages": [
                            {
                                "role": "system",
                                "content": system_prompt
                            },
                            {
                                "role": "user",
                                "content": prompt
                            }
                        ],
                        "stream": False
                    }
                )
                
                if response.status_code != 200:
                    logging.error(f"Ollama API error: {response.status_code}")
                    logging.error(f"Response content: {response.text}")
                    return f"Error: Ollama API returned status code {response.status_code}. Please check if the model '{model}' is available."
                    
                result = response.json()
                response_text = result.get('message', {}).get('content', '')
                if not response_text:
                    response_text = result.get('response', '')  # Fallback for older API versions
                response_text = response_text.strip()
                
                # Check if it's a command to execute
                if response_text.upper().startswith('EXECUTE:'):
                    command = response_text.split('EXECUTE:', 1)[1].strip()
                    logging.info(f"Executing shell command: {command}")
                    output = execute_shell_command(command)
                    return f"Command executed:\n{command}\n\nOutput:\n{output}"
                elif response_text.upper().startswith('PYTHON:'):
                    # Extract Python script between PYTHON: and END_PYTHON
                    script_parts = response_text.split('END_PYTHON', 1)
                    if len(script_parts) < 2:
                        return "Error: Python script is not properly formatted. Missing END_PYTHON marker."
                    
                    script = script_parts[0].split('PYTHON:', 1)[1].strip()
                    logging.info(f"Executing Python script:\n{script}")
                    output = execute_python_script(script)
                    return f"Python script executed:\n{script}\n\nOutput:\n{output}"
                
                logging.info(f"Ollama response received, length: {len(response_text)}")
                return response_text
            
            except Exception as e:
                logging.error(f"Error with Ollama API: {str(e)}")
                return f"Error communicating with Ollama: {str(e)}"
                
        else:
            # Use OpenAI directly instead of shell-gpt
            try:
                import openai
                
                api_key = st.session_state.get('openai_api_key', '')
                if not api_key:
                    return "Error: OpenAI API key not set. Please enter your API key in the sidebar."
                
                openai.api_key = api_key
                
                response = openai.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                return response.choices[0].message.content.strip()
                
            except Exception as e:
                error_msg = f"Error with OpenAI API: {str(e)}"
                logging.error(error_msg)
                return error_msg
    
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logging.error(error_msg)
        return error_msg
    finally:
        logging.info("Query generation complete")

def get_available_ollama_models():
    """Get available Ollama models using API"""
    if not check_ollama_running():
        logging.error("Ollama server is not running")
        return ["Error: Ollama server is not running. Please start Ollama first."]
        
    try:
        url = "http://localhost:11434/api/tags"  # Changed from /api/list to /api/tags
        logging.info(f"Attempting to connect to Ollama API at {url}")
        
        response = requests.get(url, timeout=5)
        logging.info(f"Ollama API response status: {response.status_code}")
        logging.info(f"Ollama API response: {response.text[:200]}")
        
        response.raise_for_status()
        result = response.json()
        
        # Extract model names from the response
        models = []
        if isinstance(result, dict) and 'models' in result:
            for model in result['models']:
                if isinstance(model, dict) and 'name' in model:
                    name = model['name']
                    if name:
                        models.append(name)
                        logging.info(f"Found model: {name}")
        else:
            logging.error(f"Unexpected API response structure: {result}")
            return ["Error: Unexpected API response format"]
        
        if not models:
            logging.warning("No models found in Ollama")
            return ["No models installed. Use the 'Download New Models' section to install models."]
            
        logging.info(f"Successfully found {len(models)} models: {', '.join(models)}")
        return sorted(models)
        
    except requests.exceptions.Timeout:
        logging.error("Timeout connecting to Ollama API")
        return ["Error: Connection to Ollama timed out"]
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching Ollama models: {str(e)}")
        return ["Error loading Ollama models. Please check if Ollama is running properly."]
    except json.JSONDecodeError as e:
        logging.error(f"Error parsing Ollama API response: {str(e)}")
        return ["Error: Invalid response from Ollama API"]
    except Exception as e:
        logging.error(f"Unexpected error fetching models: {str(e)}")
        return [f"Error: {str(e)}"]

def download_ollama_model(model_name):
    """Download an Ollama model"""
    try:
        url = "http://localhost:11434/api/pull"
        payload = {"name": model_name}
        
        # Check if Ollama is running first
        if not check_ollama_running():
            return False, "Error: Ollama server is not running. Please start Ollama first."
        
        with st.spinner(f"Downloading {model_name}..."):
            # Initialize progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Start the download
            response = requests.post(url, json=payload, stream=True)
            response.raise_for_status()
            
            # Process the streaming response
            total_size = 0
            downloaded_size = 0
            status_count = 0  # To track status updates
            
            for line in response.iter_lines():
                if line:
                    data = json.loads(line)
                    
                    # Handle different types of status updates
                    if 'total' in data:
                        total_size = data['total']
                    if 'completed' in data:
                        downloaded_size = data['completed']
                        if total_size > 0:
                            progress = min(downloaded_size / total_size, 1.0)
                            progress_bar.progress(progress)
                    
                    # Update status message
                    if 'status' in data:
                        status = data['status']
                        status_text.write(f"Status: {status}")
                        status_count += 1
                        
                        # Log progress for debugging
                        if status_count % 10 == 0:  # Log every 10th status update
                            logging.info(f"Download progress for {model_name}: {status}")
                    
                    # Handle errors
                    if 'error' in data:
                        error_msg = data['error']
                        logging.error(f"Error downloading {model_name}: {error_msg}")
                        return False, f"Error downloading model: {error_msg}"
            
            # Clear progress indicators
            progress_bar.empty()
            status_text.empty()
            
            # Verify the model was installed
            installed_models = get_available_ollama_models()
            if model_name.split(':')[0] in [m.split(':')[0] for m in installed_models]:
                logging.info(f"Successfully installed model: {model_name}")
                return True, f"Successfully downloaded model: {model_name}"
            else:
                logging.error(f"Model {model_name} not found after download")
                return False, f"Error: Model {model_name} was not found after download. Please try again."
            
    except requests.exceptions.ConnectionError:
        logging.error(f"Connection error while downloading {model_name}")
        return False, "Error: Could not connect to Ollama server. Please check if Ollama is running."
    except requests.exceptions.HTTPError as e:
        logging.error(f"HTTP error downloading {model_name}: {str(e)}")
        return False, f"Error downloading model: HTTP error {e.response.status_code}"
    except requests.exceptions.RequestException as e:
        logging.error(f"Error downloading model: {str(e)}")
        return False, f"Error downloading model: {str(e)}"
    except json.JSONDecodeError as e:
        logging.error(f"Error parsing response while downloading {model_name}: {str(e)}")
        return False, "Error: Invalid response from Ollama server"
    except Exception as e:
        logging.error(f"Unexpected error downloading {model_name}: {str(e)}")
        return False, f"Unexpected error: {str(e)}"

def delete_ollama_model(model_name):
    """Delete an Ollama model"""
    try:
        url = "http://localhost:11434/api/delete"
        payload = {"name": model_name}
        
        logging.info(f"Attempting to delete model {model_name} via {url}")
        response = requests.delete(url, json=payload)
        logging.info(f"Delete response status: {response.status_code}")
        logging.info(f"Delete response content: {response.text}")
        
        response.raise_for_status()
        logging.info(f"Successfully deleted model: {model_name}")
        return True, f"Successfully deleted model: {model_name}"
            
    except requests.exceptions.RequestException as e:
        error_msg = f"Error deleting model: {str(e)}"
        logging.error(error_msg)
        if hasattr(e.response, 'text'):
            logging.error(f"Error response content: {e.response.text}")
        return False, error_msg

def get_downloadable_ollama_models():
    """Get list of downloadable Ollama models"""
    try:
        # List of known available models
        # Source: https://ollama.ai/library
        available_models = [
            # Base Models
            "llama2", "mistral", "phi", "solar", "codellama", "falcon", "orca-mini", "vicuna",
            # Specialized Models
            "codellama:python", "codellama:code", "deepseek-coder", "wizard-math", "nous-hermes",
            # Multilingual Models
            "llama2-chinese", "llama2-japanese", "llama2-german", "llama2-korean", "llama2-french",
            # Vision Models
            "llava", "bakllava", "llama2-vision",
            # Chat Models
            "neural-chat", "starling-lm", "openchat", "stable-beluga", "zephyr",
            # Uncensored Models
            "llama2-uncensored", "wizard-vicuna", "dolphin-phi",
            # Latest Additions
            "mixtral", "qwen", "yi", "gemma", "phi-2", "mistral-medium", "mistral-small",
            "mixtral-instruct", "nous-hermes2", "nous-hermes2-mixtral", "neural-chat-7b-v3-1",
            "deepseek-llm", "deepseek-coder:33b", "stable-code", "stablelm2", "tinyllama",
            "llama2:13b", "llama2:70b", "mistral:7b", "mistral:7b-instruct"
        ]
        
        # Get currently installed models
        installed_models = get_available_ollama_models()
        # Filter out error messages from installed models
        installed_models = [m for m in installed_models if not m.startswith("Error")]
        
        # Create a set of installed base model names (handling versions and variants)
        installed_base_models = set()
        for model in installed_models:
            # Split on common separators (: and -)
            parts = model.replace('-', ':').split(':')
            base_name = parts[0].lower()  # Base name is always the first part
            installed_base_models.add(base_name)
            # If it's a versioned/variant model, also add the second part
            if len(parts) > 1:
                variant = f"{base_name}:{parts[1].lower()}"
                installed_base_models.add(variant)
        
        logging.info(f"Found {len(installed_base_models)} installed models: {', '.join(installed_base_models)}")
        
        # Return models that aren't already installed (comparing base names and variants)
        downloadable = []
        for model in available_models:
            model_lower = model.lower()
            base_name = model_lower.split(':')[0]
            
            # Add model if neither its base name nor full name is installed
            if base_name not in installed_base_models and model_lower not in installed_base_models:
                downloadable.append(model)
        
        downloadable.sort()
        logging.info(f"Found {len(downloadable)} downloadable models: {', '.join(downloadable)}")
        return downloadable
    except Exception as e:
        logging.error(f"Error getting downloadable models: {str(e)}")
        return []

def read_logs(num_lines=50):
    """Read the last n lines of the log file"""
    try:
        with open(log_file, 'r') as f:
            lines = f.readlines()
            return ''.join(lines[-num_lines:])
    except Exception as e:
        return f"Error reading logs: {str(e)}"

def copy_to_clipboard(text):
    """Copy text to clipboard and show success message"""
    try:
        pyperclip.copy(text)
        return True
    except Exception as e:
        logging.error(f"Failed to copy to clipboard: {str(e)}")
        return False

def display_response_with_copy(response):
    if 'response_container' not in st.session_state:
        st.session_state.response_container = response
        st.session_state.response_counter = 0
    
    # Update the response container if a new response is provided
    if response != st.session_state.response_container:
        st.session_state.response_container = response
        st.session_state.response_counter += 1
    
    col1, col2 = st.columns([0.95, 0.05])
    
    with col1:
        st.markdown(st.session_state.response_container)
    
    with col2:
        # Create a unique key for the clear button using both content hash and counter
        clear_key = f"clear_btn_{hashlib.sha256(response.encode()).hexdigest()[:8]}_{st.session_state.response_counter}"
        if st.button("🗑", help="Clear response", key=clear_key):
            st.session_state.response_container = ""
            # Add a flag to indicate the response was cleared
            st.session_state.response_cleared = True
            return True
    
    return False

def handle_user_input(user_prompt):
    """Handle user input and generate appropriate response"""
    if not user_prompt:
        st.warning("Please enter a prompt first.")
        return
    
    # Reset the cleared flag when starting a new query
    if 'response_cleared' not in st.session_state:
        st.session_state.response_cleared = False
    
    # Force a new query if the previous response was cleared
    if st.session_state.response_cleared:
        st.session_state.response_cleared = False
        st.session_state.response_container = ""
    
    with st.spinner("Generating response..."):
        logging.info("Starting new query generation")
        
        # Check if this is a web search request
        if user_prompt.lower().startswith("search:"):
            # Handle web search
            query = user_prompt[7:].strip()  # Remove "search:" prefix
            if not query:
                st.warning("Please enter a search query after 'search:'")
                return
                
            logging.info(f"Executing web search for: {query}")
            
            # First get the web search results
            with st.spinner("Searching the web..."):
                web_results = execute_web_search(query)
                
                if web_results.startswith("Error:"):
                    st.error(web_results)
                    return
            
            # Then pass the results to the AI for analysis
            with st.spinner("Analyzing search results..."):
                enhanced_prompt = f"""Analyze and summarize these web search results for the query: "{query}"

Search Results:
{web_results}

Please provide a clear and informative response that includes:
1. A summary of the key points and main ideas
2. Any relevant details, data, or specific information
3. Context and insights that might be helpful
4. If the search results appear to be incomplete or irrelevant, please mention that

Format the response in a clear, readable way using markdown formatting where appropriate."""

                response = run_sgpt_command(
                    enhanced_prompt,
                    model=st.session_state.selected_model,
                    use_ollama=st.session_state.use_ollama
                )
        else:
            response = None
            content = process_uploaded_files(uploaded_files, st.session_state.selected_model)
            
            if content:  # If we have uploaded file content
                logging.info("Processing uploaded file content")
                messages = generate_ollama_prompt(user_prompt, content, st.session_state.selected_model)
                response_placeholder = st.empty()
                full_response = ""
                
                try:
                    for chunk in query_ollama(messages, st.session_state.selected_model):
                        full_response += chunk
                        # Update the placeholder with the current response
                        with response_placeholder:
                            st.markdown(full_response)  # Just display the markdown without the clear button during streaming
                        response = full_response
                    # After streaming is complete, display the final response with the clear button
                    if response:
                        if not display_response_with_copy(response):  # Only store if not cleared
                            handle_conversation_storage(user_prompt, response, st.session_state.current_model)
                except Exception as e:
                    st.error(f"Error generating response: {str(e)}")
                    logging.error(f"Error in response generation: {str(e)}")
                
            elif user_prompt.lower().startswith('search:'):
                # Handle web search
                query = user_prompt[7:].strip()
                search_results = web_search(query)
                response = process_search_results(search_results, query, include_analysis=True)
                if response and not response.startswith("Error:"):
                    if not display_response_with_copy(response):  # Only store if not cleared
                        handle_conversation_storage(user_prompt, response, st.session_state.current_model)
            else:
                # Regular shell-gpt processing without file content
                response = run_sgpt_command(user_prompt, st.session_state.selected_model, st.session_state.use_ollama)
                if response and not response.startswith("Error:"):
                    if not display_response_with_copy(response):  # Only store if not cleared
                        handle_conversation_storage(user_prompt, response, st.session_state.current_model)
                
            # Handle errors
            if not response or response.startswith("Error:"):
                st.error(response if response else "No response generated")
            logging.info("Query generation complete")

def process_uploaded_files(uploaded_files, selected_model):
    """Process uploaded files and generate content for the prompt"""
    if not uploaded_files:
        return []
        
    content = []
    for uploaded_file in uploaded_files:
        try:
            file_type = magic.from_buffer(uploaded_file.read(), mime=True)
            uploaded_file.seek(0)  # Reset file pointer
            logging.info(f"Processing file: {uploaded_file.name} ({file_type})")

            if file_type.startswith('image/'):
                if not is_vision_model(selected_model):
                    st.warning(f"Selected model '{selected_model}' does not support image analysis. Please select a vision-capable model.")
                    continue
                    
                image_content = process_image(uploaded_file)
                if image_content:
                    content.append({"type": "image", "data": image_content})
                    logging.info(f"Successfully processed image: {uploaded_file.name}")
            elif file_type == 'application/pdf':
                text_content = extract_text_from_pdf(uploaded_file)
                if text_content:
                    # Store PDF analysis in session state
                    st.session_state.file_analysis = {
                        'file_type': 'pdf',
                        'text_content': text_content,
                        'documentation': {
                            'has_docstrings': False  # Not applicable for PDFs
                        }
                    }
                    content.append({"type": "text", "data": f"Content from PDF file '{uploaded_file.name}':\n\n{text_content}"})
                    logging.info(f"Successfully extracted text from PDF: {uploaded_file.name} (length: {len(text_content)})")
                else:
                    logging.warning(f"No text content extracted from PDF: {uploaded_file.name}")
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                text_content = extract_text_from_docx(uploaded_file)
                if text_content:
                    content.append({"type": "text", "data": f"Content from Word document '{uploaded_file.name}':\n\n{text_content}"})
                    logging.info(f"Successfully extracted text from DOCX: {uploaded_file.name} (length: {len(text_content)})")
                else:
                    logging.warning(f"No text content extracted from DOCX: {uploaded_file.name}")
            elif file_type.startswith('text/'):
                text_content = uploaded_file.read().decode('utf-8', errors='replace')
                content.append({"type": "text", "data": f"Content from text file '{uploaded_file.name}':\n\n{text_content}"})
                logging.info(f"Successfully read text file: {uploaded_file.name} (length: {len(text_content)})")
            else:
                st.warning(f"Unsupported file type: {file_type}")
                logging.warning(f"Unsupported file type: {file_type} for file: {uploaded_file.name}")
                
        except Exception as e:
            logging.error(f"Error processing file {uploaded_file.name}: {str(e)}")
            st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
    
    logging.info(f"Processed {len(content)} files successfully")
    return content

def is_vision_model(model_name):
    """Check if the model supports vision/image analysis"""
    vision_models = [
        "llama2-vision",
        "llama3.2-vision",
        "bakllava",
        "llava"
    ]
    return any(vision_model.lower() in model_name.lower() for vision_model in vision_models)

def process_image(image_file):
    """Process an image file and return base64 encoded content"""
    try:
        image_bytes = image_file.getvalue()  # Use getvalue() for StreamingUploadedFile
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        return base64_image
    except Exception as e:
        logging.error(f"Error processing image: {str(e)}")
        return None

def generate_ollama_prompt(query, content, model):
    """Generate a prompt for Ollama API including images if present"""
    messages = []
    
    # Add system message based on content type
    if any(item["type"] == "image" for item in content):
        messages.append({
            "role": "system",
            "content": "You are an AI assistant capable of understanding and analyzing images. Provide detailed, accurate descriptions and insights about any images shared."
        })
    elif any(item["type"] == "text" for item in content):
        messages.append({
            "role": "system",
            "content": "You are an AI assistant capable of analyzing documents and text content. Please provide detailed insights and analysis of the documents shared."
        })
    
    # Add content to the messages
    for item in content:
        if item["type"] == "image":
            messages.append({
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "data": item["data"]
                    }
                ]
            })
        elif item["type"] == "text":
            messages.append({
                "role": "user",
                "content": item["data"]
            })
    
    # Add the user's query with enhanced instructions
    enhanced_query = f"""Please analyze the provided document(s) with the following query: {query}

Please provide:
1. A summary of the key points and main ideas
2. Any relevant details, data, or specific information
3. Context and insights that might be helpful
4. If the search results appear to be incomplete or irrelevant, please mention that

Format the response in a clear, readable way using markdown formatting where appropriate."""

    messages.append({
        "role": "user",
        "content": enhanced_query
    })
    
    return messages

def query_ollama(messages, model):
    """Query Ollama API with support for both text and images"""
    try:
        url = "http://localhost:11434/api/generate"  
        
        # Format the prompt and images
        prompt = ""
        images = []
        
        for msg in messages:
            if isinstance(msg.get('content'), list):
                # Handle image content
                for content_item in msg['content']:
                    if content_item['type'] == 'image':
                        images.append(content_item['data'])
            else:
                # Handle text content
                prompt += msg.get('content', '') + "\n"
        
        payload = {
            "model": model,
            "prompt": prompt.strip(),
            "stream": True
        }
        
        # Add images if present
        if images:
            payload["images"] = images
            
        logging.info(f"Sending request to Ollama with model: {model}")
        logging.info(f"Prompt: {prompt[:100]}...")  
        
        response = requests.post(url, json=payload, stream=True)
        response.raise_for_status()
        
        for line in response.iter_lines():
            if line:
                try:
                    chunk = json.loads(line)
                    if 'error' in chunk:
                        raise Exception(chunk['error'])
                    if 'response' in chunk:
                        yield chunk['response']
                except json.JSONDecodeError as e:
                    logging.error(f"Error parsing chunk: {str(e)}")
                    continue
                    
    except Exception as e:
        error_msg = f"Error querying Ollama: {str(e)}"
        logging.error(error_msg)
        yield error_msg

def extract_text_from_pdf(file_bytes):
    """Extract text from PDF file"""
    try:
        # Handle Streamlit UploadedFile object
        if hasattr(file_bytes, 'read'):
            file_bytes = file_bytes.read()
            
        # For complex PDF files, use pypdf
        pdf_reader = PdfReader(BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip() if text.strip() else None
        
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {str(e)}")
        return None

def extract_text_from_docx(file_bytes_or_path):
    """Extract text from DOCX file"""
    try:
        if isinstance(file_bytes_or_path, str):
            # If input is a file path
            doc = docx.Document(file_bytes_or_path)
        else:
            # If input is bytes
            doc = docx.Document(BytesIO(file_bytes_or_path))
            
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {str(e)}")
        return None

def check_ollama_running():
    """Check if Ollama server is running"""
    try:
        logging.info("Checking if Ollama server is running...")
        response = requests.get("http://localhost:11434/api/version", timeout=5)
        logging.info(f"Ollama version check response: {response.status_code}")
        return response.status_code == 200
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to connect to Ollama: {str(e)}")
        return False

def process_search_results(results, query=None, include_analysis=True):
    """Process and format search results with optional LLM analysis"""
    try:
        formatted_results = []
        
        # Add LLM analysis if requested
        if include_analysis and query and results:
            analysis = analyze_search_results(query, results)
            formatted_results.append("🤖 AI Analysis:")
            formatted_results.append(analysis)
            formatted_results.append("\n📊 Raw Search Results:")
        
        # Format raw search results
        for idx, result in enumerate(results, 1):
            formatted_results.append(f"{idx}. {result['title']}")
            formatted_results.append(f"   {result['snippet']}")
            formatted_results.append(f"   🔗 {result['link']}")
            formatted_results.append("")  # Empty line between results
        
        return "\n".join(formatted_results)
    except Exception as e:
        logging.error(f"Error processing search results: {str(e)}")
        return "Error processing search results"

def analyze_search_results(query, results):
    """Analyze search results using the LLM and provide a summary"""
    try:
        # Prepare the content for LLM analysis
        content = f"Search Query: {query}\n\nSearch Results:\n"
        for idx, result in enumerate(results, 1):
            content += f"\n{idx}. {result['title']}\n"
            content += f"   {result['snippet']}\n"
            content += f"   Source: {result['link']}\n"

        # Prepare the prompt for the LLM
        prompt = f"""Based on the following search results, provide a concise but comprehensive summary that answers the search query. 
If the results don't directly answer the query, mention that and provide the most relevant information available.
If the results seem contradictory, point that out and explain the different viewpoints.

{content}

Please format your response in this structure:
1. Direct Answer (if available)
2. Key Points
3. Additional Context (if relevant)
"""

        if st.session_state.use_ollama:
            messages = [{"role": "user", "content": prompt}]
            model = st.session_state.selected_model or "llama2"
            
            # Make a direct API call to Ollama
            url = "http://localhost:11434/api/generate"
            data = {
                "model": model,
                "prompt": prompt,
                "stream": False  # Don't stream for analysis
            }
            
            response = requests.post(url, json=data)
            response.raise_for_status()
            response_data = response.json()
            analysis = response_data.get('response', 'Error analyzing results with Ollama')
        else:
            client = OpenAI(api_key=st.session_state.openai_api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )
            analysis = response.choices[0].message.content

        return analysis
    except Exception as e:
        logging.error(f"Error analyzing search results: {str(e)}")
        return f"Error analyzing search results: {str(e)}"

def web_search(query, num_results=10):
    """Perform a web search and return results using DuckDuckGo"""
    try:
        logging.info(f"Performing web search for: {query}")
        url = f"https://duckduckgo.com/html/?q={query}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        results = []
        
        # Find all search result elements
        for result in soup.select('.result'):
            if len(results) >= num_results:
                break
                
            title_elem = result.select_one('.result__title')
            snippet_elem = result.select_one('.result__snippet')
            link_elem = result.select_one('.result__url')
            
            if title_elem and snippet_elem and link_elem:
                title = title_elem.get_text(strip=True)
                snippet = snippet_elem.get_text(strip=True)
                link = link_elem.get_text(strip=True)
                
                results.append({
                    "title": title,
                    "snippet": snippet,
                    "link": link
                })
                logging.info(f"Found result: {title}")
        
        if not results:
            # Fallback in case no results were found
            results.append({
                "title": "No detailed results found",
                "snippet": f"Please try rephrasing your search query: {query}",
                "link": url
            })
        
        logging.info(f"Found {len(results)} search results")
        return results
        
    except Exception as e:
        logging.error(f"Error in web search: {str(e)}")
        raise Exception(f"Error executing web search: {str(e)}")

def display_conversation_history():
    """This function is now deprecated as the conversation history is handled in the main function."""
    pass  # Conversation history is now displayed in the main function

def handle_conversation_storage(prompt: str, response: str, model_name: str):
    """Store a conversation interaction in the database."""
    try:
        # Get or create conversation
        if not st.session_state.get('current_conversation_id'):
            # Create new conversation
            title = generate_conversation_title(prompt)
            st.session_state.current_conversation_id = db.create_conversation(
                model_name=model_name,
                title=title
            )
            logging.info(f"Created new conversation with ID {st.session_state.current_conversation_id}")
        
        # Store messages
        db.add_message(st.session_state.current_conversation_id, "user", prompt)
        db.add_message(st.session_state.current_conversation_id, "assistant", response)
        
        # Update session state
        conversation, messages = db.get_conversation(st.session_state.current_conversation_id)
        st.session_state.conversation_messages = messages
        
        logging.info(f"Stored conversation interaction in conversation {st.session_state.current_conversation_id}")
    except Exception as e:
        logging.error(f"Error storing conversation: {str(e)}")
        # Reset conversation ID if there was an error
        st.session_state.current_conversation_id = None

def generate_conversation_title(first_message: str, max_length: int = 50) -> str:
    """Generate a title for a conversation based on the first message."""
    # Take the first sentence or first few words
    title = first_message.split('.')[0].strip()
    if len(title) > max_length:
        title = title[:max_length-3] + "..."
    return title

def initialize_workspace_manager():
    """Initialize the workspace manager with the current directory."""
    workspace_dir = os.getcwd()
    return WorkspaceManager(workspace_dir)

def display_workspace_section():
    """Display the workspace analysis and code improvement section."""
    with st.sidebar.expander("📁 Workspace", expanded=False):
        st.markdown("""
            <style>
                .streamlit-expanderHeader {
                    font-size: 1.2rem !important;
                    font-weight: 600 !important;
                }
                /* Style for workspace content */
                .workspace-content {
                    max-width: 100%;
                    overflow-x: hidden;
                }
                /* Style for file path display */
                .file-path {
                    white-space: nowrap;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    max-width: 100%;
                    padding: 0.25rem 0;
                }
                /* Style for file buttons */
                .file-button {
                    width: 100%;
                    text-align: left;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    white-space: nowrap;
                }
                /* Style for action buttons container */
                .action-buttons {
                    display: flex;
                    justify-content: flex-end;
                    gap: 0.25rem;
                }
                /* Style for metrics cards */
                .metric-card {
                    background-color: rgba(255, 255, 255, 0.05);
                    border-radius: 0.5rem;
                    padding: 1rem;
                    margin: 0.5rem 0;
                }
                .metric-title {
                    font-size: 0.9rem;
                    color: #9BA3AF;
                    margin-bottom: 0.25rem;
                }
                .metric-value {
                    font-size: 1.2rem;
                    font-weight: 600;
                }
            </style>
            """, unsafe_allow_html=True)
            
        st.markdown('<div class="workspace-content">', unsafe_allow_html=True)
        
        # Initialize workspace manager
        workspace_manager = initialize_workspace_manager()
        
        # Tabs for different workspace features
        tab1, tab2, tab3 = st.tabs(["📁 Files", "📊 Analysis", "💡 Insights"])
        
        with tab1:
            # File Management Section
            st.markdown("### File Management")
            
            # Display workspace path
            workspace_path = workspace_manager.get_workspace_path()
            st.markdown(f'<div class="file-path">📁 {workspace_path}</div>', unsafe_allow_html=True)
            
            # File upload
            uploaded_files = st.file_uploader(
                "Add files to workspace",
                accept_multiple_files=True,
                key="workspace_files"
            )
            
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    try:
                        # Save directly to the workspace directory
                        target_path = workspace_manager.get_workspace_path() / uploaded_file.name
                        with open(target_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        show_alert(f"Added: {uploaded_file.name}", "success")
                    except Exception as e:
                        show_alert(f"Failed to add: {uploaded_file.name}", "error")
                        logging.error(f"Error adding file to workspace: {str(e)}")
            
            # List and manage existing files
            workspace_files = workspace_manager.list_workspace_files()
            if workspace_files:
                st.markdown("#### Workspace Files")
                for file_path in workspace_files:
                    # Container for each file row
                    st.markdown(f"""
                        <div style="
                            display: flex;
                            justify-content: space-between;
                            align-items: center;
                            margin-bottom: 0.5rem;
                            padding: 0.25rem;
                            border-radius: 0.25rem;
                            background-color: rgba(255, 255, 255, 0.05);
                        ">
                            <div style="
                                flex-grow: 1;
                                overflow: hidden;
                                text-overflow: ellipsis;
                                white-space: nowrap;
                                margin-right: 0.5rem;
                            ">
                                {file_path.name}
                            </div>
                        </div>
                    """, unsafe_allow_html=True)
                    
                    # Action buttons in a separate row
                    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
                    with col1:
                        if st.button("📂", key=f"open_{file_path.name}", help="Open file"):
                            if open_file(file_path):
                                show_alert(f"Opened: {file_path.name}", "success")
                            else:
                                show_alert(f"Failed to open: {file_path.name}", "error")
                    with col2:
                        if st.button("📋", key=f"copy_{file_path.name}", help="Copy path"):
                            pyperclip.copy(str(file_path))
                    with col3:
                        if st.button("✏️", key=f"edit_{file_path.name}", help="Edit file"):
                            st.session_state.editing_file = str(file_path)
                            st.session_state.show_editor = True
                            show_alert(f"Editing: {file_path.name}", "info")
                    with col4:
                        if st.button("🗑", key=f"remove_{file_path.name}", help="Delete file"):
                            if workspace_manager.remove_from_workspace(file_path.name):
                                show_alert(f"Deleted: {file_path.name}", "success")
                                st.rerun()
        
        with tab2:
            # Code Analysis Section
            st.markdown("### Code Analysis")
            
            # Analyze button for whole workspace
            if st.button("🔄 Analyze Workspace", use_container_width=True):
                with st.spinner("Analyzing workspace..."):
                    analysis = workspace_manager.analyze_codebase()
                    st.session_state.workspace_analysis = analysis
            
            # Display analysis results if available
            if 'workspace_analysis' in st.session_state:
                analysis = st.session_state.workspace_analysis
                
                # Overview metrics
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("""
                        <div class="metric-card">
                            <div class="metric-title">Total Files</div>
                            <div class="metric-value">{}</div>
                        </div>
                    """.format(analysis['total_files']), unsafe_allow_html=True)
                
                with col2:
                    st.markdown("""
                        <div class="metric-card">
                            <div class="metric-title">Languages</div>
                            <div class="metric-value">{}</div>
                        </div>
                    """.format(len(analysis['languages'])), unsafe_allow_html=True)
                
                # Detailed metrics sections using markdown
                st.markdown("### Detailed Metrics")
                st.markdown("#### File Types")
                for ext, count in analysis['file_types'].items():
                    st.markdown(f"- `{ext}`: {count} files")
                
                st.markdown("#### Languages")
                for lang in analysis['languages']:
                    st.markdown(f"- {lang}")
                
                st.markdown("#### Dependencies")
                if analysis['dependencies']:
                    for dep in analysis['dependencies']:
                        st.markdown(f"- `{dep}`")
                
                st.markdown("#### Documentation Coverage")
                doc_coverage = analysis['documentation_coverage']
                st.markdown(f"- Files with docs: {doc_coverage['files_with_docs']}")
                st.markdown(f"- Total docstrings: {doc_coverage['total_docstrings']}")
                
                st.markdown("#### Test Coverage")
                test_cov = analysis['test_coverage']
                st.markdown(f"- Test files: {test_cov['test_files']}")
                st.markdown("#### Test Directories")
                for test_dir in test_cov['test_directories']:
                    st.markdown(f"- {test_dir}")
                
                st.markdown("#### Complexity Metrics")
                complexity = analysis['complexity_metrics']
                st.markdown(f"- Average file size: {complexity['avg_file_size']:.2f} bytes")
                st.markdown(f"- Max file size: {complexity['max_file_size']} bytes")
                st.markdown("#### Largest Files")
                for file_path, size in complexity['largest_files']:
                    st.markdown(f"- `{file_path}`: {size} bytes")
        
        with tab3:
            # Code Insights Section
            st.markdown("### Code Insights")
            
            # File selection for analysis
            files = workspace_manager.list_workspace_files()
            if files:
                selected_file = st.selectbox(
                    "Select file for analysis",
                    options=[str(f) for f in files],
                    format_func=lambda x: Path(x).name
                )
                
                if st.button("🔍 Analyze File", use_container_width=True):
                    with st.spinner("Analyzing file..."):
                        suggestions = workspace_manager.suggest_code_improvements(selected_file)
                        st.session_state.file_suggestions = suggestions
                        
                        # Get detailed file context
                        context = workspace_manager.get_file_context(selected_file)
                        st.session_state.file_analysis = context
                        show_alert(f"Analyzed: {selected_file}", "info")
                
                # Display file analysis if available
                if 'file_analysis' in st.session_state:
                    context = st.session_state.file_analysis
                    
                    # Handle different file types
                    file_type = context.get('file_type', 'unknown')
                    
                    if file_type == 'pdf':
                        st.markdown("#### PDF Analysis")
                        if 'text_content' in context:
                            with st.expander("PDF Content"):
                                st.text(context['text_content'])
                            # Add PDF-specific metrics
                            text = context['text_content']
                            word_count = len(text.split())
                            char_count = len(text)
                            st.markdown(f"- Word count: {word_count}")
                            st.markdown(f"- Character count: {char_count}")
                    else:
                        # Display code metrics for source code files
                        if 'complexity_metrics' in context:
                            st.markdown("#### Code Metrics")
                            metrics = context['complexity_metrics']
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown(f"- Lines of code: {metrics['lines_of_code']}")
                                st.markdown(f"- Functions: {metrics['function_count']}")
                                st.markdown(f"- Classes: {metrics['class_count']}")
                            with col2:
                                st.markdown(f"- Comments: {metrics['comment_lines']}")
                                st.markdown(f"- Blank lines: {metrics['blank_lines']}")
                                st.markdown(f"- Imports: {metrics['import_count']}")
                    
                    # Documentation status - only show for code files
                    if file_type != 'pdf' and 'documentation' in context:
                        st.markdown("#### Documentation")
                        docs = context['documentation']
                        st.markdown(f"- Has docstrings: {'✅' if docs['has_docstrings'] else '❌'}")
                        st.markdown(f"- Docstring count: {docs['docstring_count']}")
                        st.markdown(f"- TODO items: {docs['todo_count']}")
                
                # Display improvement suggestions if available
                if 'file_suggestions' in st.session_state:
                    suggestions = st.session_state.file_suggestions
                    
                    if suggestions:
                        st.markdown("#### Suggested Improvements")
                        for suggestion in suggestions:
                            severity_color = {
                                'error': '🔴',
                                'warning': '🟡',
                                'suggestion': '🔵',
                                'info': '⚪'
                            }.get(suggestion['severity'], '⚪')
                            
                            st.markdown(f"{severity_color} **{suggestion['type'].title()}**: {suggestion['message']}")
                    else:
                        st.info("No suggestions found. The file follows good practices! 🎉")
        
        st.markdown('</div>', unsafe_allow_html=True)

def open_file(file_path):
    """Open a file using the default system application."""
    try:
        if sys.platform == "darwin":  # macOS
            subprocess.run(["open", str(file_path)])
        elif sys.platform == "win32":  # Windows
            os.startfile(str(file_path))
        else:  # Linux
            subprocess.run(["xdg-open", str(file_path)])
        logging.info(f"Opened file: {file_path}")
        return True
    except Exception as e:
        logging.error(f"Error opening file: {str(e)}")
        return False

def detect_language(file_path: str) -> str:
    """Detect the programming language based on file extension."""
    ext_to_language = {
        '.py': 'python',
        '.js': 'javascript',
        '.jsx': 'javascript',
        '.ts': 'typescript',
        '.tsx': 'typescript',
        '.html': 'html',
        '.css': 'css',
        '.scss': 'scss',
        '.json': 'json',
        '.md': 'markdown',
        '.sql': 'sql',
        '.sh': 'bash',
        '.bash': 'bash',
        '.zsh': 'bash',
        '.cpp': 'cpp',
        '.c': 'c',
        '.h': 'c',
        '.hpp': 'cpp',
        '.java': 'java',
        '.rb': 'ruby',
        '.php': 'php',
        '.go': 'go',
        '.rs': 'rust',
        '.swift': 'swift',
        '.kt': 'kotlin',
        '.r': 'r',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.xml': 'xml',
        '.txt': 'text'
    }
    ext = Path(file_path).suffix.lower()
    return ext_to_language.get(ext, 'text')

def display_file_editor():
    """Display the file editor interface with syntax highlighting."""
    if not hasattr(st.session_state, 'show_editor') or not st.session_state.show_editor:
        return

    try:
        file_path = st.session_state.editing_file
        if not file_path or not os.path.exists(file_path):
            show_alert("No file selected or file does not exist", "error")
            st.session_state.show_editor = False
            return

        with open(file_path, 'r') as f:
            content = f.read()

        # Initialize editor state if needed
        if 'editor_mode' not in st.session_state:
            st.session_state.editor_mode = 'view'
        if 'editor_content' not in st.session_state:
            st.session_state.editor_content = content
        if 'editor_language' not in st.session_state:
            st.session_state.editor_language = detect_language(file_path)

        st.markdown("### File Editor")
        
        # Top controls row
        col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 2, 1])
        
        with col1:
            if st.session_state.editor_mode == 'edit':
                if st.button("Save", key="save_btn"):
                    try:
                        with open(file_path, 'w') as f:
                            f.write(st.session_state.editor_content)
                        show_alert("File saved successfully!", "success")
                    except Exception as e:
                        show_alert(f"Failed to save file: {str(e)}", "error")
        
        with col2:
            if st.button("Close", key="close_btn"):
                st.session_state.show_editor = False
                st.session_state.editing_file = None
                st.session_state.editor_mode = 'view'
                if 'editor_content' in st.session_state:
                    del st.session_state.editor_content
                show_alert("Editor closed", "info")
                st.rerun()
        
        with col3:
            # Toggle between view and edit modes
            new_mode = st.button(
                "Edit" if st.session_state.editor_mode == 'view' else "View",
                key="toggle_mode"
            )
            if new_mode:
                st.session_state.editor_mode = 'view' if st.session_state.editor_mode == 'edit' else 'edit'
                st.rerun()
        
        with col4:
            st.text(f"Editing: {Path(file_path).name}")
        
        with col5:
            # Language selector with proper label
            st.session_state.editor_language = st.selectbox(
                "Language",
                options=['python', 'javascript', 'typescript', 'html', 'css', 'json', 'markdown', 
                        'sql', 'bash', 'cpp', 'java', 'ruby', 'php', 'go', 'rust', 'swift', 
                        'kotlin', 'r', 'yaml', 'xml', 'text'],
                index=['python', 'javascript', 'typescript', 'html', 'css', 'json', 'markdown', 
                      'sql', 'bash', 'cpp', 'java', 'ruby', 'php', 'go', 'rust', 'swift', 
                      'kotlin', 'r', 'yaml', 'xml', 'text'].index(st.session_state.editor_language),
                key="language_select",
                label_visibility="collapsed"
            )

        # Add a separator
        st.markdown("---")

        # Editor content
        if st.session_state.editor_mode == 'edit':
            st.session_state.editor_content = st.text_area(
                "File Contents",
                value=content,  # Use the content directly here
                height=400,
                key="file_editor"
            )
        else:
            # View mode with syntax highlighting
            st.code(
                content,  # Use the content directly here
                language=st.session_state.editor_language,
                line_numbers=True
            )

    except Exception as e:
        logging.error(f"Error in file editor: {str(e)}")
        show_alert(f"Error in file editor: {str(e)}", "error")
        st.session_state.show_editor = False
        st.session_state.editing_file = None
        if 'editor_content' in st.session_state:
            del st.session_state.editor_content

def show_alert(message: str, type: str = "info"):
    """Display a custom alert banner at the top of the UI.
    
    Args:
        message: The message to display
        type: One of "success", "error", "warning", "info"
    """
    # Generate a unique key for this alert
    alert_id = hashlib.md5(f"{message}{time.time()}".encode()).hexdigest()
    
    if 'alerts' not in st.session_state:
        st.session_state.alerts = []
    
    # Add the alert to session state if it's not already there
    if not any(alert['id'] == alert_id for alert in st.session_state.alerts):
        st.session_state.alerts.append({
            'id': alert_id,
            'message': message,
            'type': type
        })

def display_alerts():
    """Display all active alerts at the top of the UI."""
    if 'alerts' not in st.session_state:
        st.session_state.alerts = []
    
    for alert in st.session_state.alerts[:]:  # Create a copy to allow removal during iteration
        # Custom styling based on alert type
        if alert['type'] == 'success':
            bgcolor = '#D4EDDA'
            bordercolor = '#C3E6CB'
            textcolor = '#155724'
            icon = '✅'
        elif alert['type'] == 'error':
            bgcolor = '#F8D7DA'
            bordercolor = '#F5C6CB'
            textcolor = '#721C24'
            icon = '❌'
        elif alert['type'] == 'warning':
            bgcolor = '#FFF3CD'
            bordercolor = '#FFEEBA'
            textcolor = '#856404'
            icon = '⚠️'
        else:  # info
            bgcolor = '#D1ECF1'
            bordercolor = '#BEE5EB'
            textcolor = '#0C5460'
            icon = 'ℹ️'

        # Create columns for alert content and close button
        col1, col2 = st.columns([20, 1])
        
        with col1:
            st.markdown(f"""
                <div style="
                    padding: 1rem;
                    margin-bottom: 1rem;
                    border: 1px solid {bordercolor};
                    border-radius: 0.25rem;
                    background-color: {bgcolor};
                    color: {textcolor};
                    display: flex;
                    align-items: center;
                ">
                    <span style="margin-right: 0.5rem;">{icon}</span>
                    <span>{alert['message']}</span>
                </div>
            """, unsafe_allow_html=True)
        
        with col2:
            if st.button("×", key=f"close_alert_{alert['id']}", help="Dismiss alert"):
                st.session_state.alerts.remove(alert)
                st.rerun()

def get_cache_key(prompt: str, model: str) -> str:
    """Generate a cache key from prompt and model."""
    return hashlib.md5(f"{prompt}:{model}".encode()).hexdigest()

@lru_cache(maxsize=100)
def get_cached_response(cache_key: str) -> Optional[str]:
    """Get cached response if available."""
    try:
        cache_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".cache", f"{cache_key}.json")
        if os.path.exists(cache_file):
            with open(cache_file, 'r') as f:
                data = json.load(f)
                # Cache entries expire after 24 hours
                if time.time() - data['timestamp'] < 86400:
                    return data['response']
    except Exception as e:
        logging.error(f"Cache read error: {str(e)}")
    return None

def save_to_cache(cache_key: str, response: str) -> None:
    """Save response to cache."""
    try:
        cache_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".cache")
        os.makedirs(cache_dir, exist_ok=True)
        cache_file = os.path.join(cache_dir, f"{cache_key}.json")
        with open(cache_file, 'w') as f:
            json.dump({
                'response': response,
                'timestamp': time.time()
            }, f)
    except Exception as e:
        logging.error(f"Cache write error: {str(e)}")

def generate_query(prompt, model, use_ollama=True, temperature=0.7, max_tokens=150):
    """Generate a response using the specified model"""
    try:
        # Ensure we have a valid model
        if not model:
            model = "mistral" if use_ollama else "gpt-3.5-turbo"
            logging.warning(f"No model specified, using default: {model}")
        
        # Check cache first
        cache_key = get_cache_key(prompt, model)
        cached_response = get_cached_response(cache_key)
        if cached_response:
            logging.info("Using cached response")
            return cached_response
        
        # Adjust parameters based on prompt length
        prompt_length = len(prompt)
        if prompt_length > 1000:
            temperature = max(0.3, temperature - 0.2)  # Reduce temperature for long prompts
            max_tokens = min(max_tokens, 4000 - prompt_length)  # Adjust max_tokens to fit context window
        
        if use_ollama:
            if not check_ollama_running():
                return "Error: Ollama server is not running. Please start it first."
            
            response = query_ollama([{"role": "user", "content": prompt}], model)
        else:
            client = OpenAI(api_key=st.session_state.openai_api_key)
            
            # Stream the response
            placeholder = st.empty()
            full_response = []
            
            for chunk in client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature,
                max_tokens=max_tokens,
                stream=True
            ):
                if chunk.choices[0].delta.content is not None:
                    full_response.append(chunk.choices[0].delta.content)
                    # Update the placeholder with the current response
                    with placeholder:
                        st.markdown(full_response)  # Just display the markdown without the clear button during streaming
                    response = full_response
            # After streaming is complete, display the final response with the clear button
            if response:
                if not display_response_with_copy(response):  # Only store if not cleared
                    handle_conversation_storage(prompt, response, model)
            
            placeholder.empty()

        # Cache the response
        save_to_cache(cache_key, response)
        
        # Store conversation in database
        handle_conversation_storage(prompt, response, model)
        
        return response
    except Exception as e:
        logging.error(f"Error generating query: {str(e)}")
        return f"Error: {str(e)}"

def open_user_guide():
    """Open the user guide in the default web browser."""
    guide_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "user_guide.html")
    if os.path.exists(guide_path):
        import webbrowser
        webbrowser.open('file://' + guide_path)
    else:
        st.error("User guide not found. Please ensure the static/user_guide.html file exists.")

def main():
    """Main application function"""
    # Initialize session state for model selection if not already set
    if 'selected_model' not in st.session_state:
        default_model = load_default_model()
        st.session_state.selected_model = default_model if default_model else "gpt-3.5-turbo"
    
    if 'use_ollama' not in st.session_state:
        st.session_state.use_ollama = True
        
    if 'email_manager' not in st.session_state:
        st.session_state.email_manager = EmailManager()
    
    # Initialize all session state variables
    if 'openai_api_key' not in st.session_state:
        st.session_state.openai_api_key = load_api_key()
    if 'current_model' not in st.session_state:
        default_model = load_default_model()
        if default_model:
            st.session_state.current_model = default_model
        else:
            # Set default model based on use_ollama
            st.session_state.current_model = "mistral" if st.session_state.use_ollama else "gpt-3.5-turbo"
    if 'available_models' not in st.session_state:
        st.session_state.available_models = []
    if 'selected_content' not in st.session_state:
        st.session_state.selected_content = ""
    if 'delete_clicked' not in st.session_state:
        st.session_state.delete_clicked = False
    if 'preview_mode' not in st.session_state:
        st.session_state.preview_mode = False
    if 'response' not in st.session_state:
        st.session_state.response = ""
    if 'template_prompt' not in st.session_state:
        st.session_state.template_prompt = ""
    if 'working_directory' not in st.session_state:
        st.session_state.working_directory = os.getcwd()
    if 'directory_history' not in st.session_state:
        st.session_state.directory_history = [os.getcwd()]
    if 'current_dir_index' not in st.session_state:
        st.session_state.current_dir_index = 0
    if 'selected_file' not in st.session_state:
        st.session_state.selected_file = None
    if 'current_conversation_id' not in st.session_state:
        st.session_state.current_conversation_id = None
    if 'conversation_messages' not in st.session_state:
        st.session_state.conversation_messages = []
    if 'show_editor' not in st.session_state:
        st.session_state.show_editor = False
    if 'editing_file' not in st.session_state:
        st.session_state.editing_file = None
    if 'editor_content' not in st.session_state:
        st.session_state.editor_content = ""

    # Main interface styling and layout
    st.markdown("""
        <style>
            @font-face {
                font-family: 'Calendas Plus';
                src: url('static/fonts/calendas_plus-webfont.woff2') format('woff2'),
                     url('static/fonts/calendas_plus-webfont.woff') format('woff');
                font-weight: normal;
                font-style: normal;
            }
            
            /* Base theme colors and variables */
            :root {
                --bg-dark: #262730;
                --bg-medium: #2E3648;
                --bg-light: #363E54;
                --accent: #1E88E5;
                --text-light: #FAFAFA;
                --text-dark: #262730;
                --sidebar-width: 22.5rem;
            }

            /* Apply font family to all elements */
            * {
                font-family: 'Calendas Plus', -apple-system, BlinkMacSystemFont, sans-serif;
            }

            /* Main container styling */
            .main .block-container {
                padding: 2rem 1rem;
                max-width: none;
                width: 100%;
            }

            /* Ensure sidebar stays anchored to left */
            [data-testid="stSidebar"] {
                background-color: var(--bg-dark);
                position: fixed;
                left: 0;
                top: 0;
                height: 100vh;
                width: var(--sidebar-width) !important;
            }

            [data-testid="stSidebarContent"] {
                background-color: var(--bg-dark);
            }

            /* Adjust main content to account for fixed sidebar */
            .main .block-container {
                margin-left: var(--sidebar-width);
                width: calc(100% - var(--sidebar-width));
            }

            /* Make columns responsive */
            [data-testid="column"] {
                width: 100% !important;
                flex: 1 1 auto !important;
                min-width: 400px;
            }

            /* Ensure middle spacing column stays proportional */
            [data-testid="column"]:nth-of-type(2) {
                flex: 0 0 5% !important;
                min-width: 40px;
            }

            /* Container for all columns */
            [data-testid="column-container"] {
                width: 100% !important;
                max-width: none !important;
                gap: 2rem;
            }

            /* Navigation buttons styling */
            .nav-buttons-container {
                display: flex;
                justify-content: center;
                align-items: center;
                gap: 1rem;
                margin-bottom: 1rem;
            }
            
            .nav-buttons-container .stButton > button {
                width: 80px;
                height: 42px;
                padding: 0.5rem;
                display: flex;
                align-items: center;
                justify-content: center;
            }

            /* Button styling */
            .stButton > button {
                background-color: var(--bg-medium);
                color: var(--text-light);
                border: 1px solid var(--accent);
                border-radius: 4px;
                transition: all 0.3s ease;
            }

            .stButton > button:hover {
                background-color: var(--accent);
                color: var(--text-light);
                border-color: var(--accent);
            }

            /* Text color overrides */
            .stMarkdown, .stText {
                color: var(--text-light) !important;
            }

            /* Ensure content fills width */
            .element-container, .stTextArea, textarea {
                width: 100% !important;
            }
            
            /* Hide deploy button */
            .st-emotion-cache-1wbqy5l {
                display: none !important;
            }
        </style>
    """, unsafe_allow_html=True)
    
    # Add custom CSS to reduce top padding in sidebar
    st.markdown("""
        <style>
        [data-testid="stSidebarNav"] {
            padding-top: 0rem;
        }
        [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h1 {
            margin-top: 0rem;
            padding-top: 0.5rem;
        }
        </style>
    """, unsafe_allow_html=True)

    # Sidebar configuration
    with st.sidebar:
        # Add title with custom styling
        st.markdown("""
            <div style='text-align: center; border-bottom: 1px solid rgba(250, 250, 250, 0.2); margin-bottom: 25px; padding-bottom: 10px;'>
                <h1 style='font-size: 1.8em;'>AI Workstation</h1>
            </div>
        """, unsafe_allow_html=True)
        
        # Model Configuration in collapsible section
        with st.expander("🤖 Model Configuration", expanded=True):
            # Initialize use_ollama in session state if not present
            if 'use_ollama' not in st.session_state:
                st.session_state.use_ollama = True
                
            use_ollama = st.checkbox(
                "Use Ollama",
                value=st.session_state.use_ollama,
                key="use_ollama_checkbox",
                on_change=lambda: setattr(st.session_state, 'use_ollama', not st.session_state.use_ollama)
            )

            if use_ollama:
                available_models = get_available_ollama_models()
                downloadable_models = get_downloadable_ollama_models()
                
                if available_models or downloadable_models:
                    # Combine installed and downloadable models
                    model_options = []
                    
                    # Add installed models
                    for model in available_models:
                        if not model.startswith("Error"):
                            model_options.append(f"{model} (installed)")
                    
                    # Add downloadable models
                    for model in downloadable_models:
                        if model not in [m.split(" (")[0] for m in model_options]:
                            model_options.append(f"{model} (available)")
                    
                    # Sort the combined list
                    model_options.sort(key=lambda x: x.split(" (")[0])
                    
                    if not model_options:
                        st.warning("No Ollama models available. Please install Ollama and download at least one model.")
                        st.session_state.selected_model = None
                        return
                    
                    # Find current model in options
                    current_model = st.session_state.selected_model
                    current_model_with_status = next(
                        (opt for opt in model_options if opt.startswith(current_model + " ")),
                        model_options[0] if model_options else None
                    )
                    
                    # Model selection with status preservation
                    selected = st.selectbox(
                        "Select Ollama Model",
                        options=model_options,
                        index=model_options.index(current_model_with_status) if current_model_with_status else 0,
                        key="ollama_model_select"
                    )

                    # Extract model name and update session state
                    model_name = selected.split(" (")[0]
                    if st.session_state.selected_model != model_name:
                        st.session_state.selected_model = model_name

                    # Show model status
                    default_model = load_default_model()
                    if model_name == default_model:
                        st.markdown("🌟 Default Model")
                    else:
                        st.markdown("✓ Active Model")

                    # Download button for available models
                    if selected.endswith("(available)"):
                        if st.button(f"Download {model_name}", key="download_model"):
                            with st.spinner(f"Downloading {model_name}..."):
                                success, message = download_ollama_model(model_name)
                                if success:
                                    show_alert(message, "success")
                                    st.rerun()
                                else:
                                    show_alert(message, "error")

                    # Set as default button for installed models
                    if selected.endswith("(installed)"):
                        if st.button("Set as Default", key="set_default_ollama"):
                            save_default_model(model_name)
                            show_alert(f"Saved {model_name} as default model", "success")
                            st.rerun()
            else:
                # OpenAI model selection
                st.markdown("##### 🧠 OpenAI Models")
                col1, col2 = st.columns([3, 1])
                with col1:
                    previous_model = st.session_state.selected_model
                    st.session_state.selected_model = st.selectbox(
                        "Select Model",
                        options=["gpt-3.5-turbo", "gpt-4"],
                        index=["gpt-3.5-turbo", "gpt-4"].index(previous_model) if previous_model in ["gpt-3.5-turbo", "gpt-4"] else 0,
                        key="openai_model_select"
                    )
                
                with col2:
                    # Show model status
                    default_model = load_default_model()
                    if st.session_state.selected_model == default_model:
                        st.markdown("🌟 Default Model")
                    else:
                        st.markdown("✓ Active Model")
            
                # Set as default button
                if st.button("Set as Default", key="set_default_openai"):
                    save_default_model(st.session_state.selected_model)
                    show_alert(f"Saved {st.session_state.selected_model} as default model", "success")
                    st.rerun()
    
    # Add workspace section to sidebar
    display_workspace_section()
    display_file_editor()

    # Conversation History Section in Sidebar
    with st.sidebar:
        with st.expander("💬 Conversation History", expanded=False):
            st.markdown("""
                <style>
                    .streamlit-expanderHeader {
                        font-size: 1.2rem !important;
                        font-weight: 600 !important;
                    }
                </style>
                """, unsafe_allow_html=True)
            
            search_query = st.text_input("🔍 Search Conversations", key="search_history")
            
            try:
                # Get conversations based on search query
                if search_query:
                    conversations = db.search_conversations(search_query)
                else:
                    conversations = db.list_conversations()
                
                if conversations:
                    for conv in conversations:
                        # Create a container for each conversation
                        conv_container = st.container()
                        with conv_container:
                            # Header with title and date
                            st.markdown(f"""
                            ### 📝 {conv['title']}
                            **Model:** {conv['model_name']}  
                            **Date:** {conv['start_time']}
                            """)
                            
                            # Display messages
                            for msg in conv['messages']:
                                role_emoji = "👤" if msg['role'] == 'user' else "🤖"
                                msg_content = msg['content']
                                msg_key = f"msg_{conv['id']}_{msg['id']}"
                                
                                # Initialize expansion state in session state if not present
                                if msg_key not in st.session_state:
                                    st.session_state[msg_key] = False
                                
                                # Show first 100 characters with option to show more
                                if len(msg_content) > 100:
                                    # Create message container
                                    msg_container = st.container()
                                    
                                    # Toggle button with current state
                                    if st.button(
                                        "🔄 " + ("Show less" if st.session_state[msg_key] else "Show more"), 
                                        key=f"toggle_{msg_key}"
                                    ):
                                        st.session_state[msg_key] = not st.session_state[msg_key]
                                        st.rerun()
                                    
                                    # Display message content
                                    with msg_container:
                                        if st.session_state[msg_key]:
                                            st.markdown(f"{role_emoji} **{msg['role'].title()}:** {msg_content}")
                                        else:
                                            st.markdown(f"{role_emoji} **{msg['role'].title()}:** {msg_content[:100]}...")
                                else:
                                    st.markdown(f"{role_emoji} **{msg['role'].title()}:** {msg_content}")
                            
                            # Action buttons for the entire conversation
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.button("📋 Copy", key=f"copy_{conv['id']}", help="Copy entire conversation"):
                                    full_text = "\n\n".join([f"{msg['role'].title()}: {msg['content']}" 
                                                           for msg in conv['messages']])
                                    pyperclip.copy(full_text)
                                    show_alert("Conversation copied to clipboard", "success")
                            
                            with col2:
                                if st.button("🗑️ Delete", key=f"delete_{conv['id']}", help="Delete entire conversation"):
                                    if db.delete_conversation(conv['id']):
                                        # Reset conversation ID if we're deleting the current conversation
                                        if st.session_state.current_conversation_id == conv['id']:
                                            st.session_state.current_conversation_id = None
                                        show_alert("Conversation deleted", "success")
                                        time.sleep(0.5)  # Brief pause for feedback
                                        st.rerun()
                                    else:
                                        show_alert("Failed to delete conversation", "error")
                            
                            # Add a separator between conversations
                            st.markdown("---")
                else:
                    st.info("No conversations found")
                    if search_query:
                        st.markdown("*Try a different search term*")
            except Exception as e:
                st.error(f"Error loading conversations: {str(e)}")
                logging.error(f"Error in conversation history: {str(e)}")
        
        # Debug section at bottom of sidebar
        st.markdown("---")
        with st.expander("🛠️ Debug", expanded=False):
            # Initialize debug toggle in session state if not present
            if 'show_debug' not in st.session_state:
                st.session_state.show_debug = False
            
            # Toggle button for debug info
            if st.button("Toggle Debug Info"):
                st.session_state.show_debug = not st.session_state.show_debug
                st.rerun()
            
            # Show debug info if enabled
            if st.session_state.show_debug:
                st.markdown("#### Session State")
                st.json(st.session_state)
                
                st.markdown("#### Environment")
                st.json({
                    "working_directory": os.getcwd(),
                    "python_version": sys.version,
                    "streamlit_version": st.__version__,
                })
                
                st.markdown("#### Recent Logs")
                logs = read_logs(50)  # Get last 50 lines of logs
                if logs:
                    st.code(logs)
                else:
                    st.info("No logs available")
        
        # User guide button
        if st.button("📚 User Guide", help="Open the comprehensive user guide in a new window"):
            open_user_guide()
    
    # Set up the Streamlit interface with responsive columns
    container = st.container()
    with container:
        # Create columns with flexible widths
        left_col, middle_space, right_col = st.columns([47.5, 5, 47.5])
    
    with left_col:
        st.markdown('<div style="text-align: center;">', unsafe_allow_html=True)
        st.title("AI Agent")
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Add spacing
        st.markdown("<div style='padding: 1em'></div>", unsafe_allow_html=True)
        
        # Main prompt input
        default_prompt = st.session_state.get('template_prompt', '')
        user_prompt = st.text_area(
            label="Enter your prompt, or search query using 'search:'",
            value=default_prompt if default_prompt else "",
            height=150,
            placeholder="Enter your prompt here...",
            key="user_prompt"
        )
        
        # File upload section below text area
        uploaded_files = st.file_uploader(
            "Optional: Upload a context file for analysis",
            type=["txt", "png", "jpg", "jpeg", "pdf", "docx"],
            help="Upload a document or image to analyze alongside your prompt",
            accept_multiple_files=True
        )
        
        # Generate Response button
        if st.button("Generate Response", use_container_width=True):
            if user_prompt:
                response = None
                content = process_uploaded_files(uploaded_files, st.session_state.selected_model)
                
                if content:  # If we have uploaded file content
                    logging.info("Processing uploaded file content")
                    messages = generate_ollama_prompt(user_prompt, content, st.session_state.selected_model)
                    response_placeholder = st.empty()
                    full_response = ""
                    
                    try:
                        for chunk in query_ollama(messages, st.session_state.selected_model):
                            full_response += chunk
                            # Update the placeholder with the current response
                            with response_placeholder:
                                st.markdown(full_response)  # Just display the markdown without the clear button during streaming
                            response = full_response
                        # After streaming is complete, display the final response with the clear button
                        if response:
                            if not display_response_with_copy(response):  # Only store if not cleared
                                handle_conversation_storage(user_prompt, response, st.session_state.current_model)
                    except Exception as e:
                        st.error(f"Error generating response: {str(e)}")
                        logging.error(f"Error in response generation: {str(e)}")
                
                elif user_prompt.lower().startswith('search:'):
                    # Handle web search
                    query = user_prompt[7:].strip()
                    search_results = web_search(query)
                    response = process_search_results(search_results, query, include_analysis=True)
                    if response and not response.startswith("Error:"):
                        if not display_response_with_copy(response):  # Only store if not cleared
                            handle_conversation_storage(user_prompt, response, st.session_state.current_model)
                else:
                    # Regular shell-gpt processing without file content
                    response = run_sgpt_command(user_prompt, st.session_state.selected_model, st.session_state.use_ollama)
                    if response and not response.startswith("Error:"):
                        if not display_response_with_copy(response):  # Only store if not cleared
                            handle_conversation_storage(user_prompt, response, st.session_state.current_model)
                
                # Handle errors
                if not response or response.startswith("Error:"):
                    st.error(response if response else "No response generated")
                logging.info("Query generation complete")
        
        # Clear the template prompt after use
        if default_prompt and user_prompt != default_prompt:
            st.session_state.template_prompt = ""
        
        # Preview area
        if st.session_state.preview_mode and st.session_state.selected_content:
            st.subheader("File Preview")
            try:
                preview_container = st.container()
                with preview_container:
                    st.text_area(
                        label="File Contents",
                        value=st.session_state.selected_content,
                        height=400,
                        key="preview",
                        disabled=True
                    )
            except Exception as e:
                st.error(f"Error displaying preview: {str(e)}")
                logging.error(f"Preview error: {str(e)}")
        
        # Display response in a container with a clear button
        if st.session_state.response:
            response_container = st.container()
            col1, col2 = response_container.columns([20, 1])
            
            with col1:
                st.markdown(st.session_state.response)
            
            with col2:
                if st.button("❌", key="clear_response", help="Clear response"):
                    st.session_state.response = ""
                    st.rerun()
        
        # Copy button
        if st.session_state.response:
            if st.button("📋 Copy to Clipboard"):
                pyperclip.copy(st.session_state.response)
                st.toast("Response copied to clipboard!", icon="✅")
    
    with right_col:
        st.markdown('<div style="text-align: center;">', unsafe_allow_html=True)
        st.title("Files")
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Directory navigation with better styling
        st.markdown('<div class="nav-buttons-container">', unsafe_allow_html=True)
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("⬆️ Up", key="up_dir"):
                parent_dir = str(pathlib.Path(st.session_state.working_directory).parent)
                if parent_dir != st.session_state.working_directory:  # Prevent going above root
                    st.session_state.working_directory = parent_dir
                    if st.session_state.working_directory not in st.session_state.directory_history:
                        st.session_state.directory_history.append(st.session_state.working_directory)
                        st.session_state.current_dir_index = len(st.session_state.directory_history) - 1
        
        with col2:
            if st.button("⬇️ Down", key="down_dir"):
                # Show a warning if no subdirectories exist
                current_subdirs = [d for d in pathlib.Path(st.session_state.working_directory).iterdir() if d.is_dir()]
                if not current_subdirs:
                    st.warning("No subdirectories available")
                else:
                    # Get the first subdirectory
                    next_dir = str(current_subdirs[0])
                    st.session_state.working_directory = next_dir
                    if next_dir not in st.session_state.directory_history:
                        st.session_state.directory_history.append(next_dir)
                        st.session_state.current_dir_index = len(st.session_state.directory_history) - 1
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Current directory name
        st.markdown("### Current Location:")
        st.code(st.session_state.working_directory, language="bash")
        
        # Show full path in an expander
        with st.expander("Click to Copy Path", expanded=False):
            if st.button(st.session_state.working_directory):
                # Add to clipboard
                pyperclip.copy(st.session_state.working_directory)
                st.success("Path copied to clipboard!")
        
        # List directory contents with improved styling
        try:
            items = os.listdir(st.session_state.working_directory)
            items.sort(key=lambda x: (not os.path.isdir(os.path.join(st.session_state.working_directory, x)), x.lower()))
            
            with st.expander("### Contents", expanded=False):
                for item in items:
                    full_path = os.path.join(st.session_state.working_directory, item)
                    if os.path.isdir(full_path):
                        if st.button(f"📁 {item}", key=f"dir_{item}"):
                            st.session_state.working_directory = full_path
                            st.session_state.selected_file = None
                            st.session_state.selected_content = None
                            st.session_state.preview_expanded = False
                    else:
                        if st.button(f"📄 {item}", key=f"file_{item}"):
                            try:
                                if os.path.getsize(full_path) < 1000000:  # Only preview files < 1MB
                                    with open(full_path, 'r') as f:
                                        st.session_state.selected_file = item
                                        st.session_state.selected_content = f.read()
                                        st.session_state.preview_expanded = True
                            except:
                                st.warning(f"Cannot preview {item}")
        except Exception as e:
            st.error(f"Error accessing directory: {str(e)}")
        
        # File preview section
        if st.session_state.selected_file and st.session_state.selected_content:
            with st.expander(f"Preview: {st.session_state.selected_file}", expanded=st.session_state.preview_expanded):
                st.text_area(
                    label="File Contents",
                    value=st.session_state.selected_content,
                    height=400,
                    key="preview"
                )

    # Email Section
    email_col = st.container()
    with email_col:
        st.subheader("📧 Email Interface")
        
        # Check if nmail is installed and configured
        if not st.session_state.email_manager.is_nmail_installed():
            st.warning("nmail is not installed. Click below to install it.")
            if st.button("Install nmail"):
                success, message = st.session_state.email_manager.install_nmail()
                if success:
                    st.success(message)
                    st.experimental_rerun()
                else:
                    st.error(message)
        elif not st.session_state.email_manager.is_nmail_configured():
            st.warning("nmail needs to be configured with your email service.")
            
            # Email service selection
            services = st.session_state.email_manager.get_supported_services()
            selected_service = st.selectbox(
                "Select your email service:",
                services,
                index=0,
                help="Choose your email provider to start the setup process"
            )
            
            if selected_service == 'gmail':
                st.info("""
                ### Gmail Setup Instructions
                
                You'll need to prepare the following:
                
                1. Your Gmail address
                2. Either:
                   - An App Password (recommended for accounts with 2-factor authentication)
                   - Enable "Less secure app access" (only if you don't use 2-factor authentication)
                
                The setup process will:
                1. Guide you through creating an App Password or enabling less secure app access
                2. Configure nmail with your Gmail settings
                3. Test the connection to ensure everything works
                
                Click "Start Setup" when you're ready to begin.
                """)
            elif selected_service == 'gmail-oauth2':
                st.info("""
                ### Gmail OAuth2 Setup
                
                This method is recommended as it's more secure and doesn't require enabling less secure app access.
                
                You'll need:
                1. Your Gmail address
                2. Access to a web browser for OAuth2 authentication
                
                Click "Start Setup" to begin the OAuth2 authentication process.
                """)
            else:
                st.info(f"""
                ### {selected_service.title()} Setup
                
                The setup wizard will guide you through configuring your {selected_service.title()} account.
                
                Please have your email credentials ready before starting.
                """)
            
            setup_col1, setup_col2 = st.columns([1, 1])
            with setup_col1:
                if st.button("Start Setup", key="start_setup"):
                    success, message = st.session_state.email_manager.setup_email_service(selected_service)
                    if success:
                        st.info(message + "\n\nPlease complete the setup in the terminal window.")
                    else:
                        st.error(message)
            
            with setup_col2:
                if st.button("Check Configuration", key="check_config"):
                    if st.session_state.email_manager.is_nmail_configured():
                        st.success("Email configuration successful! You can now use the email features.")
                        st.experimental_rerun()
                    else:
                        st.error("Configuration not complete. Please finish the setup in the terminal window.")
        else:
            # Email interface
            email_tabs = st.tabs(["💬 Natural Language", "📝 Compose", "📨 Open nmail", "🔍 Debug"])
            
            # Natural Language tab
            with email_tabs[0]:
                nl_command = st.text_area("Enter your email command", 
                    placeholder="Example: send email to user@example.com subject Meeting tomorrow with body Let's meet at 2 PM")
                if st.button("Send", key="nl_send"):
                    if nl_command:
                        success, message = st.session_state.email_manager.process_natural_language_command(nl_command)
                        if success:
                            st.success(message)
                        else:
                            st.error(f"Error sending email: {message}")
                            # Show the logs
                            with st.expander("Show Debug Information"):
                                st.code(read_logs(20))
                            
            # Compose tab
            with email_tabs[1]:
                to_addr = st.text_input("To:", placeholder="recipient@example.com")
                subject = st.text_input("Subject:", placeholder="Enter subject")
                body = st.text_area("Message:", placeholder="Enter your message")
                
                if st.button("Send", key="compose_send"):
                    if all([to_addr, subject, body]):
                        success, message = st.session_state.email_manager.send_email(to_addr, subject, body)
                        if success:
                            st.success(message)
                        else:
                            st.error(f"Error sending email: {message}")
                            # Show the logs
                            with st.expander("Show Debug Information"):
                                st.code(read_logs(20))
                    else:
                        st.warning("Please fill in all fields")
                        
            # Open nmail tab
            with email_tabs[2]:
                st.write("Launch the full nmail interface in a new window")
                if st.button("Open nmail"):
                    success, message = st.session_state.email_manager.launch_nmail()
                    if success:
                        st.info("nmail launched in a new window")
                    else:
                        st.error(message)
                        
            # Debug tab
            with email_tabs[3]:
                st.markdown("### Email Configuration Debug")
                
                # Check nmail configuration
                st.subheader("nmail Configuration Status")
                config_file = Path.home() / '.config' / 'nmail' / 'main.conf'
                if config_file.exists():
                    with st.expander("Show Configuration File"):
                        try:
                            with open(config_file, 'r') as f:
                                st.code(f.read())
                        except Exception as e:
                            st.error(f"Error reading config file: {str(e)}")
                else:
                    st.warning("Configuration file not found")
                
                # Show recent logs
                st.subheader("Recent Logs")
                logs = read_logs(50)
                if logs:
                    st.code(logs)
                else:
                    st.info("No logs available")
                
                # Test nmail command
                st.subheader("Test nmail Command")
                if st.button("Test nmail"):
                    try:
                        result = subprocess.run(['nmail', '--version'], 
                                             capture_output=True, 
                                             text=True)
                        st.code(f"Exit code: {result.returncode}\nOutput: {result.stdout}\nError: {result.stderr}")
                    except Exception as e:
                        st.error(f"Error running nmail: {str(e)}")

if __name__ == "__main__":
    # Set page config (must be first Streamlit command)
    st.set_page_config(
        page_title="AI Workstation",
        page_icon="🤖",
        layout="wide"
    )
    
    # Initialize session state
    if 'openai_api_key' not in st.session_state:
        st.session_state.openai_api_key = load_api_key()
    if 'use_ollama' not in st.session_state:
        st.session_state.use_ollama = True
        
    if 'email_manager' not in st.session_state:
        st.session_state.email_manager = EmailManager()
    
    # Initialize session state variables for conversations
    if 'current_conversation_id' not in st.session_state:
        st.session_state.current_conversation_id = None
    if 'conversation_messages' not in st.session_state:
        st.session_state.conversation_messages = []
    
    # Initialize session state variables for file editor
    if 'show_editor' not in st.session_state:
        st.session_state.show_editor = False
    if 'editing_file' not in st.session_state:
        st.session_state.editing_file = None
    if 'editor_content' not in st.session_state:
        st.session_state.editor_content = ""

    display_alerts()
    main()